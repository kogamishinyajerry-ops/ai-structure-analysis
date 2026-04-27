"""Layer-4 report CLI tests — RFC-001 §3.

End-to-end runs of the thin CLI driver against the GS-001 .frd
fixture. Exit codes and stdout/stderr behaviour are part of the
contract — engineers script around them.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as _ReadDocument

from app.services.report.cli import build_parser, main


GS001_FRD = (
    Path(__file__).resolve().parents[2]
    / "golden_samples"
    / "GS-001"
    / "gs001_result.frd"
)


# --- argparse surface ----------------------------------------------------


def test_parser_requires_frd_and_kind() -> None:
    parser = build_parser()
    with pytest.raises(SystemExit):
        parser.parse_args([])


def test_parser_rejects_unknown_kind() -> None:
    parser = build_parser()
    with pytest.raises(SystemExit):
        parser.parse_args(["--frd", "x.frd", "--kind", "not-a-kind"])


def test_parser_defaults_validate_template_on() -> None:
    parser = build_parser()
    ns = parser.parse_args(["--frd", "x.frd", "--kind", "static"])
    assert ns.validate_template is True


def test_parser_can_disable_template_validation() -> None:
    parser = build_parser()
    ns = parser.parse_args(
        ["--frd", "x.frd", "--kind", "static", "--no-validate-template"]
    )
    assert ns.validate_template is False


# --- happy-path runs -----------------------------------------------------


@pytest.fixture
def gs001() -> Path:
    if not GS001_FRD.exists():
        pytest.skip(f"GS-001 .frd missing at {GS001_FRD}")
    return GS001_FRD


def test_static_run_writes_valid_docx_and_prints_summary(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "static.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "static",
            "--output",
            str(out),
        ]
    )
    assert rc == 0
    assert out.exists()
    captured = capsys.readouterr()
    assert "wrote" in captured.out
    assert "template=equipment_foundation_static" in captured.out
    # Round-trip: the .docx is real and contains EV-DISP-MAX / EV-VM-MAX.
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    assert "EV-DISP-MAX" in text
    assert "EV-VM-MAX" in text


def test_lifting_lug_run(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "lug.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "lifting-lug",
            "--output",
            str(out),
        ]
    )
    assert rc == 0
    assert out.exists()
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    assert "EV-LUG-DISP-MAX" in text
    assert "EV-LUG-VM-MAX" in text
    captured = capsys.readouterr()
    assert "template=lifting_lug" in captured.out


def test_pressure_vessel_run(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "pv.docx"
    # GS-001 is a small bracket model, not a pressure vessel — but the
    # CLI is template-agnostic about the underlying physics; the
    # producer just reports max σ along the SCL we ask for. Pick 3
    # arbitrary node ids that we know exist (GS-001 has 44 nodes).
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "pressure-vessel",
            "--output",
            str(out),
            "--scl-nodes",
            "1,2,3,4,5",
            "--scl-distances",
            "0.0, 0.5, 1.0, 1.5, 2.0",
        ]
    )
    assert rc == 0
    assert out.exists()
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    assert "EV-PM" in text
    assert "EV-PM-PB" in text
    assert "EV-MAX-VM-SCL" in text
    captured = capsys.readouterr()
    assert "template=pressure_vessel_local_stress" in captured.out


def test_pressure_vessel_requires_scl_args(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "pv.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "pressure-vessel",
                "--output",
                str(out),
            ]
        )
    # Engineers script around exit code 2 for argparse-style refusals;
    # ``SystemExit("string")`` would silently degrade to 1.
    assert excinfo.value.code == 2
    captured = capsys.readouterr()
    assert "scl-nodes" in captured.err
    assert not out.exists()


@pytest.mark.parametrize(
    "scl_nodes,scl_distances,expected_fragment",
    [
        # Trailing comma → empty field, must reject (silent drop is a
        # silent normalization the CLI does not promise).
        ("1,2,3,", "0.0,0.5,1.0", "empty field"),
        ("1,2,3", "0.0,0.5,1.0,", "empty field"),
        # Repeated comma.
        ("1,,3", "0.0,0.5,1.0", "empty field"),
        # NaN must reject — uniform-spacing check would otherwise
        # report a domain refusal (exit 3) for an input-shape problem.
        ("1,2,3", "0.0,nan,1.0", "non-finite"),
        ("1,2,3", "0.0,0.5,inf", "non-finite"),
        # Length mismatch is an input-shape error → exit 2.
        ("1,2,3,4", "0.0,0.5,1.0", "must have equal length"),
    ],
)
def test_pressure_vessel_malformed_scl_returns_2(
    gs001: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
    scl_nodes: str,
    scl_distances: str,
    expected_fragment: str,
) -> None:
    out = tmp_path / "pv.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "pressure-vessel",
                "--output",
                str(out),
                "--scl-nodes",
                scl_nodes,
                "--scl-distances",
                scl_distances,
            ]
        )
    assert excinfo.value.code == 2
    captured = capsys.readouterr()
    assert expected_fragment in captured.err
    assert not out.exists()


# --- error paths ---------------------------------------------------------


def test_missing_frd_returns_exit_code_2(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    nonexistent = tmp_path / "does_not_exist.frd"
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(nonexistent),
            "--kind",
            "static",
            "--output",
            str(out),
        ]
    )
    assert rc == 2
    assert not out.exists()
    captured = capsys.readouterr()
    assert "is not a file" in captured.err


def test_pressure_vessel_unknown_node_id_returns_3(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "pv.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "pressure-vessel",
            "--output",
            str(out),
            "--scl-nodes",
            "999999,999998,999997",  # not in GS-001
            "--scl-distances",
            "0.0,0.5,1.0",
        ]
    )
    assert rc == 3
    assert not out.exists()
    captured = capsys.readouterr()
    assert "not present in reader" in captured.err


def test_pressure_vessel_non_uniform_distances_returns_3(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "pv.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "pressure-vessel",
            "--output",
            str(out),
            "--scl-nodes",
            "1,2,3,4,5",
            "--scl-distances",
            "0.0, 0.1, 0.4, 1.0, 2.0",  # non-uniform
        ]
    )
    assert rc == 3
    assert not out.exists()
    captured = capsys.readouterr()
    assert "uniformly-spaced" in captured.err


def test_pressure_vessel_resample_accepts_non_uniform_input(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """The headline engineer flow: same non-uniform --scl-distances
    input that test_pressure_vessel_non_uniform_distances_returns_3
    rejects, but with --resample 21 added, must succeed end-to-end."""
    out = tmp_path / "pv.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "pressure-vessel",
            "--output",
            str(out),
            "--scl-nodes",
            "1,2,3,4,5",
            "--scl-distances",
            "0.0, 0.1, 0.4, 1.0, 2.0",  # non-uniform
            "--resample",
            "21",
        ]
    )
    assert rc == 0
    assert out.exists()
    captured = capsys.readouterr()
    assert "template=pressure_vessel_local_stress" in captured.out


def test_pressure_vessel_resample_below_2_returns_2(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """--resample N must reject N<2 with exit code 2 (input-shape
    error). The Layer-3 helper rejects same; the CLI surfaces it
    early as an argparse-class issue."""
    out = tmp_path / "pv.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "pressure-vessel",
                "--output",
                str(out),
                "--scl-nodes",
                "1,2,3,4,5",
                "--scl-distances",
                "0.0, 0.5, 1.0, 1.5, 2.0",
                "--resample",
                "1",
            ]
        )
    assert excinfo.value.code == 2
    captured = capsys.readouterr()
    assert "--resample must be >= 2" in captured.err
    assert not out.exists()


def test_pressure_vessel_resample_zero_returns_2(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Zero is a special case worth a separate probe — argparse parses
    it as a valid int but our domain check refuses."""
    out = tmp_path / "pv.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "pressure-vessel",
                "--output",
                str(out),
                "--scl-nodes",
                "1,2,3,4,5",
                "--scl-distances",
                "0.0, 0.5, 1.0, 1.5, 2.0",
                "--resample",
                "0",
            ]
        )
    assert excinfo.value.code == 2


def test_resample_with_static_kind_returns_2(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """--resample is only meaningful for --kind=pressure-vessel.
    Silently dropping it for static / lifting-lug would be a
    foot-gun (the engineer never learns their flag was ignored).
    Reject with exit 2 + clear stderr."""
    out = tmp_path / "static.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "static",
                "--output",
                str(out),
                "--resample",
                "21",
            ]
        )
    assert excinfo.value.code == 2
    captured = capsys.readouterr()
    assert "--resample is only meaningful with --kind=pressure-vessel" in captured.err
    assert not out.exists()


def test_resample_with_lifting_lug_kind_returns_2(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """Same foot-gun guard for --kind=lifting-lug."""
    out = tmp_path / "lug.docx"
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001),
                "--kind",
                "lifting-lug",
                "--output",
                str(out),
                "--resample",
                "21",
            ]
        )
    assert excinfo.value.code == 2
    captured = capsys.readouterr()
    assert "--resample is only meaningful with --kind=pressure-vessel" in captured.err
    assert not out.exists()


def test_parser_resample_default_is_none() -> None:
    """Without --resample, args.resample is None — the producer's
    opt-in default. Locked so the CLI default doesn't accidentally
    flip to a non-None value."""
    parser = build_parser()
    ns = parser.parse_args(
        ["--frd", "x.frd", "--kind", "pressure-vessel"]
    )
    assert ns.resample is None


def test_no_validate_template_skips_pre_flight(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    """When --no-validate-template is passed, exporter still applies
    ADR-012 cited-evidence checks but skips the template title /
    citation-count contract."""
    out = tmp_path / "static.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-validate-template",
        ]
    )
    assert rc == 0
    assert out.exists()


# --- identity-default derivation -----------------------------------------


def test_identity_defaults_derived_from_frd_stem(
    gs001: Path, tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "static.docx"
    rc = main(
        [
            "--frd",
            str(gs001),
            "--kind",
            "static",
            "--output",
            str(out),
        ]
    )
    assert rc == 0
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    # Title page surfaces project/report/bundle IDs (task_id lives on
    # the EvidenceBundle object, not the ReportSpec, so it shows up
    # via the bundle-id reference rather than as its own line).
    assert "P-gs001_result" in text
    assert "R-gs001_result" in text
    assert "B-gs001_result" in text
