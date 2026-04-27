"""Layer-4 DOCX exporter tests — RFC-001 §3 + ADR-012.

The exporter wraps python-docx; tests round-trip through the same
library to assert the rendered file's content. We do NOT pin a
particular XML byte-layout — only the user-visible structure (heading
text, paragraph text, evidence resolution).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document as _ReadDocument

from app.adapters.calculix import CalculiXReader
from app.core.types import UnitSystem
from app.models import (
    EvidenceBundle,
    EvidenceItem,
    EvidenceType,
    ReportSection,
    ReportSpec,
    SimulationEvidence,
)
from app.services.report.draft import generate_static_strength_summary
from app.services.report.exporter import (
    ExportError,
    export_docx,
    find_cited_evidence_ids,
)


GS001_FRD = (
    Path(__file__).resolve().parents[2] / "golden_samples" / "GS-001" / "gs001_result.frd"
)


# --- citation extraction ---------------------------------------------------


def test_find_cited_evidence_ids_picks_up_inline_refs() -> None:
    secs = [
        ReportSection(
            title="t",
            level=1,
            content="- 最大位移: **0.123 mm** *(EV-DISP-MAX)*",
        ),
    ]
    assert find_cited_evidence_ids(secs) == {"EV-DISP-MAX"}


def test_find_cited_evidence_ids_walks_subsections() -> None:
    secs = [
        ReportSection(
            title="parent",
            level=1,
            content="See EV-PARENT below.",
            subsections=[
                ReportSection(title="child", level=2, content="(EV-CHILD)"),
            ],
        )
    ]
    assert find_cited_evidence_ids(secs) == {"EV-PARENT", "EV-CHILD"}


def test_find_cited_evidence_ids_ignores_lowercase_or_partial() -> None:
    """Citation regex is uppercase-only; lowercase 'ev-foo' is prose."""
    secs = [
        ReportSection(
            title="t",
            level=1,
            content="lowercase ev-foo and partial EV are not citations",
        ),
    ]
    # 'EV' alone has no trailing chars matching the pattern so it falls
    # through. Empty set means the regex is correctly strict.
    assert find_cited_evidence_ids(secs) == set()


def test_find_cited_evidence_ids_handles_none_content() -> None:
    secs = [ReportSection(title="t", level=1, content=None)]
    assert find_cited_evidence_ids(secs) == set()


# --- ADR-012 hard refusals -------------------------------------------------


def _make_minimal_pair(
    *,
    bundle_evidence_ids: list[str],
    section_content: str,
    bundle_id: str = "B",
    report_bundle_id: str | None = None,
) -> tuple[ReportSpec, EvidenceBundle]:
    items = [
        EvidenceItem(
            evidence_id=eid,
            evidence_type=EvidenceType.SIMULATION,
            title=f"Synthetic {eid}",
            data=SimulationEvidence(value=1.0, unit="mm", location="node 1"),
            source="synthetic",
        )
        for eid in bundle_evidence_ids
    ]
    bundle = EvidenceBundle(bundle_id=bundle_id, task_id="T", title="b")
    for it in items:
        bundle.add_evidence(it)
    report = ReportSpec(
        report_id="R",
        project_id="P",
        title="t",
        template_id="x",
        sections=[ReportSection(title="s", level=1, content=section_content)],
        generated_at=datetime(2026, 4, 27, 0, 0, 0),
        evidence_bundle_id=report_bundle_id if report_bundle_id is not None else bundle_id,
    )
    return report, bundle


def test_export_refuses_unresolved_citation(tmp_path: Path) -> None:
    """ADR-012: section cites EV-FOO but bundle holds only EV-BAR → refuse."""
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-BAR"],
        section_content="this references *(EV-FOO)*, which doesn't exist",
    )
    with pytest.raises(ExportError, match=r"EV-FOO.*do not resolve"):
        export_docx(report, bundle, output_path=tmp_path / "out.docx")


def test_export_refuses_bundle_id_mismatch(tmp_path: Path) -> None:
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content="cites EV-OK",
        bundle_id="B-actual",
        report_bundle_id="B-stale",
    )
    with pytest.raises(ExportError, match="bundle linkage broken"):
        export_docx(report, bundle, output_path=tmp_path / "out.docx")


def test_export_refuses_missing_output_dir(tmp_path: Path) -> None:
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content="cites EV-OK",
    )
    target = tmp_path / "does_not_exist" / "out.docx"
    with pytest.raises(ExportError, match="parent .* does not exist"):
        export_docx(report, bundle, output_path=target)


def test_export_refuses_uncited_content_line(tmp_path: Path) -> None:
    """Codex R1 HIGH: ADR-012 requires every claim to cite an EV-*.
    A bare prose line without any citation is exactly the failure
    mode the rule was created to prevent (RFC-001 §2.4 rule 1)."""
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content=(
            "- 最大位移: **0.5 mm** *(EV-OK)*\n"
            "Max displacement is 0.5 mm."  # uncited prose claim — must refuse
        ),
    )
    with pytest.raises(ExportError, match="uncited content line"):
        export_docx(report, bundle, output_path=tmp_path / "out.docx")


def test_export_allows_blank_lines_inside_content(tmp_path: Path) -> None:
    """Blank lines between citations are paragraph breaks, not claims —
    they must NOT trip the uncited-content check."""
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-A", "EV-B"],
        section_content="- *(EV-A)* line one\n\n- *(EV-B)* line two",
    )
    out = export_docx(report, bundle, output_path=tmp_path / "out.docx")
    assert out.exists()


def test_export_refuses_when_parent_is_a_file(tmp_path: Path) -> None:
    """Codex R1 MEDIUM: a regular file at the parent slot would otherwise
    leak NotADirectoryError from doc.save. Surface ExportError instead."""
    parent_collision = tmp_path / "i_am_a_file"
    parent_collision.write_bytes(b"this is a file, not a directory")
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content="cites EV-OK",
    )
    with pytest.raises(ExportError, match="not a directory"):
        export_docx(report, bundle, output_path=parent_collision / "out.docx")


def test_export_returns_resolved_absolute_path(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Codex R1 MEDIUM: relative paths must be resolved on the way out
    so callers always receive an absolute path back."""
    monkeypatch.chdir(tmp_path)
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content="cites EV-OK",
    )
    relative = Path("relative-out.docx")
    assert not relative.is_absolute()
    out = export_docx(report, bundle, output_path=relative)
    assert out.is_absolute()
    assert out.exists()
    # Resolved against tmp_path cwd.
    assert out.parent == tmp_path.resolve()


def test_export_handles_deep_subsection_tree(tmp_path: Path) -> None:
    """Codex R1 MEDIUM: schema doesn't cap subsection nesting depth.
    1000 levels of subsection must not blow Python recursion limit."""
    leaf = ReportSection(title="leaf", level=3, content="cites EV-DEEP")
    cur = leaf
    for i in range(999):
        cur = ReportSection(
            title=f"L{i}", level=3, content=None, subsections=[cur]
        )
    items = [
        EvidenceItem(
            evidence_id="EV-DEEP",
            evidence_type=EvidenceType.SIMULATION,
            title="t",
            data=SimulationEvidence(value=1.0, unit="mm", location="node 1"),
            source="synthetic",
        )
    ]
    bundle = EvidenceBundle(bundle_id="B", task_id="T", title="b")
    for it in items:
        bundle.add_evidence(it)
    report = ReportSpec(
        report_id="R", project_id="P", title="t", template_id="x",
        sections=[cur], evidence_bundle_id="B",
    )
    out = export_docx(report, bundle, output_path=tmp_path / "deep.docx")
    assert out.exists()


# --- happy-path render -----------------------------------------------------


def test_export_writes_valid_docx_with_expected_text(tmp_path: Path) -> None:
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-DISP-MAX"],
        section_content=(
            "- 最大位移: **0.5 mm** @ node 1  *(EV-DISP-MAX)*\n"
            "Plain trailing line referencing EV-DISP-MAX."
        ),
    )
    out = export_docx(report, bundle, output_path=tmp_path / "out.docx")
    assert out.exists()
    assert out.stat().st_size > 0
    doc = _ReadDocument(str(out))

    all_text = "\n".join(p.text for p in doc.paragraphs)
    # Title page items
    assert "Report ID: R" in all_text
    assert "Project ID: P" in all_text
    assert "Evidence bundle: B" in all_text
    # Section content rendered
    assert "0.5 mm" in all_text
    assert "EV-DISP-MAX" in all_text
    # Appendix rendered
    assert "证据清单" in all_text or "Evidence audit trail" in all_text
    assert "Synthetic EV-DISP-MAX" in all_text


def test_export_heading_levels_map_to_docx(tmp_path: Path) -> None:
    """level=1/2/3 → Heading 1/2/3 in DOCX."""
    sections = [
        ReportSection(
            title="L1", level=1, content="cites EV-ONE",
            subsections=[
                ReportSection(
                    title="L2", level=2, content="cites EV-TWO",
                    subsections=[
                        ReportSection(title="L3", level=3, content="cites EV-THREE"),
                    ],
                ),
            ],
        )
    ]
    items = [
        EvidenceItem(
            evidence_id=f"EV-{n}",
            evidence_type=EvidenceType.SIMULATION,
            title=f"t{n}",
            data=SimulationEvidence(value=1.0, unit="mm", location="node 1"),
            source="synthetic",
        )
        for n in ("ONE", "TWO", "THREE")
    ]
    bundle = EvidenceBundle(bundle_id="B", task_id="T", title="b")
    for it in items:
        bundle.add_evidence(it)
    report = ReportSpec(
        report_id="R", project_id="P", title="t", template_id="x",
        sections=sections, evidence_bundle_id="B",
    )

    out = export_docx(report, bundle, output_path=tmp_path / "out.docx")
    doc = _ReadDocument(str(out))

    # Map heading text → style name. python-docx style names are
    # "Heading 0" (Title) ... "Heading 9".
    style_by_text = {
        p.text: (p.style.name if p.style is not None else "")
        for p in doc.paragraphs
        if p.style is not None and p.style.name.startswith("Heading")
    }
    assert style_by_text.get("L1") == "Heading 1"
    assert style_by_text.get("L2") == "Heading 2"
    assert style_by_text.get("L3") == "Heading 3"


def test_export_appendix_lists_every_evidence_item(tmp_path: Path) -> None:
    items = [
        EvidenceItem(
            evidence_id=f"EV-{i}",
            evidence_type=EvidenceType.SIMULATION,
            title=f"Item {i}",
            data=SimulationEvidence(value=float(i), unit="MPa", location=f"node {i}"),
            source="synthetic",
            source_file=f"/tmp/ghost-{i}.frd",
        )
        for i in (1, 2, 3)
    ]
    bundle = EvidenceBundle(bundle_id="B", task_id="T", title="b")
    for it in items:
        bundle.add_evidence(it)
    report = ReportSpec(
        report_id="R", project_id="P", title="t", template_id="x",
        sections=[
            ReportSection(
                title="s", level=1,
                content="cites *(EV-1)*, *(EV-2)*, *(EV-3)*",
            )
        ],
        evidence_bundle_id="B",
    )
    out = export_docx(report, bundle, output_path=tmp_path / "out.docx")
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    for i in (1, 2, 3):
        assert f"EV-{i}" in text
        assert f"Item {i}" in text
        # source_file appears in the sub-paragraph
        assert f"ghost-{i}.frd" in text


# --- end-to-end via draft generator ---------------------------------------


@pytest.fixture
def gs001_reader() -> CalculiXReader:
    if not GS001_FRD.exists():
        pytest.skip(f"GS-001 .frd missing at {GS001_FRD}")
    return CalculiXReader(GS001_FRD, unit_system=UnitSystem.SI_MM)


def test_export_e2e_calculix_to_docx(
    gs001_reader: CalculiXReader, tmp_path: Path
) -> None:
    """Full L1→L3→L4 stack: CalculiX adapter → draft generator → DOCX."""
    report, bundle = generate_static_strength_summary(
        gs001_reader,
        project_id="P-001", task_id="GS-001",
        report_id="R-001", bundle_id="B-001",
    )
    out = export_docx(report, bundle, output_path=tmp_path / "gs001.docx")
    assert out.exists()
    text = "\n".join(p.text for p in _ReadDocument(str(out)).paragraphs)
    assert "EV-DISP-MAX" in text
    assert "EV-VM-MAX" in text
    assert "Static-strength summary" in text
    # Unit pinned to SI_MM by the reader
    assert "mm" in text
    assert "MPa" in text


def test_export_overwrites_existing_file(tmp_path: Path) -> None:
    """Overwrite policy: caller manages backups; exporter clobbers."""
    out_path = tmp_path / "out.docx"
    out_path.write_bytes(b"stale content not a real docx")
    report, bundle = _make_minimal_pair(
        bundle_evidence_ids=["EV-OK"],
        section_content="cites EV-OK",
    )
    export_docx(report, bundle, output_path=out_path)
    # If we got here without error and the file is now a real docx, OK.
    doc = _ReadDocument(str(out_path))
    assert any("EV-OK" in p.text for p in doc.paragraphs)
