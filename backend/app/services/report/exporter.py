"""Layer-4 DOCX exporter — RFC-001 §3 + ADR-012.

Materialises a signed-ready ``.docx`` from a ``(ReportSpec,
EvidenceBundle)`` pair produced by :func:`~app.services.report.draft.
generate_static_strength_summary`.

ADR-012 hard refusal:
  * The exporter MUST refuse to emit DOCX for a draft whose
    section-content cites an ``evidence_id`` that doesn't resolve in
    the linked bundle.
  * The bundle linkage itself is checked: ``report.evidence_bundle_id``
    must equal ``bundle.bundle_id``.

Citation convention (matches what ``draft.py`` emits):
  Inline references appear as ``EV-<UPPER-WITH-DASHES>`` somewhere in
  the section content (typically wrapped as ``*(EV-DISP-MAX)*``). The
  exporter scans content with a strict regex (see ``_CITATION_RE``)
  and treats every match as a claim that must resolve.

What this module does NOT do (deferred):
  * Standards-citation lookup (GB / ASME) — RFC §6.4 W4+.
  * Pretty Markdown → DOCX rendering. We do *minimum-viable*
    formatting: bold ``**...**``, italics ``*...*``, leading-dash
    bullets, and headings via section ``level``. Tables, images,
    code blocks are out of scope.
  * Signature blocks. The MVP wedge says the engineer signs the
    PDF/DOCX themselves through the company's review channel
    (RFC §2.3).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING, Dict, Iterable, List, Optional, Set

from docx import Document
from docx.document import Document as _DocxDocument

from app.models import (
    EvidenceBundle,
    EvidenceItem,
    ReportSection,
    ReportSpec,
    SimulationEvidence,
)

if TYPE_CHECKING:
    from app.core.types import Material
    from app.services.report.templates import TemplateSpec


__all__ = ["export_docx", "ExportError", "find_cited_evidence_ids"]


CITATION_RE = re.compile(r"\bEV-[A-Z0-9][A-Z0-9_\-]*\b")
"""Citation pattern. ``EV-`` prefix followed by ``[A-Z0-9]`` then any
``[A-Z0-9_-]``. Word-boundary anchored so a stray ``EV-FOO`` inside a
URL won't accidentally match a leading hyphen.

The pattern intentionally rejects lowercase / mixed-case to avoid
false-positive collisions with prose. Authors who deviate must
introduce the new convention via RFC.

This is the single source of truth for the citation regex; downstream
validators (e.g. :mod:`templates`) import it directly to avoid drift
between rendering and validation."""

_CITATION_RE = CITATION_RE  # backwards-compat alias for in-module callers


class ExportError(ValueError):
    """Raised when ADR-012 invariants are violated at export time."""


def find_cited_evidence_ids(sections: Iterable[ReportSection]) -> Set[str]:
    """Walk a section tree and return the set of evidence_ids cited in
    any section's ``content`` (matched against :data:`_CITATION_RE`)."""
    cited: Set[str] = set()
    stack: List[ReportSection] = list(sections)
    while stack:
        sec = stack.pop()
        if sec.content:
            cited.update(_CITATION_RE.findall(sec.content))
        stack.extend(sec.subsections)
    return cited


def _check_every_content_line_cites_evidence(
    sections: Iterable[ReportSection],
) -> None:
    """RFC-001 §2.4 rule 1: every non-blank content line must reference
    an ``EV-*`` evidence_id. A bare prose paragraph that *describes* a
    quantity without citing its provenance is exactly the failure mode
    ADR-012 was created to prevent.

    Section *titles* are exempt (they're navigational, not claims) and
    sections with ``content=None`` are exempt (heading-only chapters).
    """
    stack: List[ReportSection] = list(sections)
    while stack:
        sec = stack.pop()
        if sec.content:
            for raw in sec.content.splitlines():
                line = raw.strip()
                if not line:
                    continue
                if not _CITATION_RE.search(line):
                    raise ExportError(
                        f"section {sec.title!r} has uncited content line: "
                        f"{line!r} — every claim must reference an EV-* "
                        "evidence_id (RFC-001 §2.4 rule 1, ADR-012)."
                    )
        stack.extend(sec.subsections)


def _format_evidence_value(item: EvidenceItem) -> str:
    """Human-readable one-liner for an evidence item, used in the
    appendix. Currently only ``SimulationEvidence`` is rendered with
    value/unit/location detail; reference / analytical evidence get a
    title-only entry until those payloads land in the MVP."""
    data = item.data
    if isinstance(data, SimulationEvidence):
        loc = f" @ {data.location}" if data.location else ""
        return f"{data.value:.6g} {data.unit}{loc}"
    return data.kind


_EMPHASIS_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _emit_runs(para: object, line: str) -> None:
    """Tokenize ``line`` into runs honouring ``**bold**`` and ``*italic*``.

    Minimum-viable: a single linear pass through :data:`_EMPHASIS_RE`.
    Doesn't handle escapes or nested emphasis — and doesn't need to,
    the draft generator emits flat one-line bullets.
    """
    if not line:
        return
    for tok in _EMPHASIS_RE.split(line):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) >= 4:
            run = para.add_run(tok[2:-2])  # type: ignore[attr-defined]
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2:
            run = para.add_run(tok[1:-1])  # type: ignore[attr-defined]
            run.italic = True
        else:
            para.add_run(tok)  # type: ignore[attr-defined]


def _render_section_content(doc: _DocxDocument, content: str) -> None:
    """Render a section's content body. Each line becomes a paragraph;
    leading ``- `` is treated as a List Bullet style."""
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            _emit_runs(para, line[2:])
        else:
            para = doc.add_paragraph()
            _emit_runs(para, line)


def _render_section_tree(
    doc: _DocxDocument, sections: Iterable[ReportSection]
) -> None:
    """Iteratively render the ReportSection tree (pre-order DFS).

    ``ReportSection.level`` is 1..3 (validated by the schema), but
    subsection nesting depth is unbounded. An iterative walk avoids a
    Python recursion-limit crash on deep trees (Codex R1 finding).
    DOCX heading levels max at 9 in python-docx; we clamp.
    """
    stack: List[ReportSection] = list(reversed(list(sections)))
    while stack:
        sec = stack.pop()
        doc.add_heading(sec.title, level=min(9, sec.level))
        if sec.content:
            _render_section_content(doc, sec.content)
        for child in reversed(sec.subsections):
            stack.append(child)


def _render_evidence_appendix(
    doc: _DocxDocument, bundle: EvidenceBundle
) -> None:
    """ADR-012 audit trail. Every evidence item gets one bulleted
    paragraph: ``- [EV-ID] Title — value unit @ location  (source: X)``.
    Empty bundles are still rendered with an explicit "no evidence"
    line — though :func:`export_docx` will never reach this branch
    (the draft generator already refuses zero-evidence reports)."""
    doc.add_heading("附录: 证据清单 / Evidence audit trail", level=1)
    if not bundle.evidence_items:
        doc.add_paragraph("(no evidence items in this bundle)")
        return
    for item in bundle.evidence_items:
        line = (
            f"- [{item.evidence_id}] {item.title} — "
            f"{_format_evidence_value(item)}  (source: {item.source})"
        )
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(line[2:])  # strip leading "- " (style adds bullet)
        if item.source_file:
            sub = doc.add_paragraph(
                f"    file: {item.source_file}", style="Intense Quote"
            )
            for run in sub.runs:
                run.italic = True


def _render_material_section(
    doc: _DocxDocument, material: "Material"
) -> None:
    """W6a / ADR-019 § 材料属性 section.

    Renders a 2-column table: 字段 (Chinese label) | 值 (numeric +
    unit). The table is the visible auditing surface — every value
    here is what the engineer compares against the design code at
    sign-off time.

    ``is_user_supplied=True`` adds a red-ink-equivalent caveat below
    the table (RFC-001 §2.4 rule 4 — low-confidence outputs must
    surface a flag the engineer must explicitly clear). The DOCX
    style uses Intense Quote + italic so the caveat stands apart from
    body text.

    si-mm convention: E / σ_y / σ_u in MPa. Density (if present) in
    tonne/mm³. Older / non-si-mm material cards come with
    ``unit_system`` other than SI_MM and the caller is responsible
    for unit conversion (out of scope for W6a).
    """
    doc.add_heading("材料属性 / Material Properties", level=1)

    # Two-column "label / value" table mirroring the GB / ASME row
    # layout common in 化工/电力 design-institute reports.
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    rows: list[tuple[str, str]] = [
        ("牌号 / Grade", f"{material.code_grade} ({material.code_standard})"),
        ("名称 / Name", material.name),
        ("弹性模量 E", f"{material.youngs_modulus:.6g} MPa"),
        ("泊松比 ν", f"{material.poissons_ratio:.3f}"),
        (
            "屈服强度 σ_y",
            f"{material.yield_strength:.6g} MPa",
        ),
        (
            "抗拉强度 σ_u",
            f"{material.ultimate_strength:.6g} MPa",
        ),
    ]
    if material.density is not None:
        rows.append(("密度 ρ", f"{material.density:.6g} tonne/mm³"))
    rows.append(("标准引用 / Citation", material.source_citation))

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    if material.is_user_supplied:
        caveat = doc.add_paragraph(
            "⚠ [需工程师确认] 本材料数据由工程师自录入，请与设计依据规范交叉核对后再签字。",
            style="Intense Quote",
        )
        for run in caveat.runs:
            run.italic = True


def _render_figures_appendix(
    doc: _DocxDocument, figures: "Dict[str, Path]"
) -> None:
    """Add a 'Figures' appendix that embeds each PNG with a short caption.

    ``figures`` maps figure-name → PNG path. Names are dictionary keys
    (insertion-ordered in modern Python), so the caller controls the
    order. Each figure is added as ``add_picture`` followed by a
    small italic caption with the figure name. Width is fixed at 6 in.
    (the typical Word page width less margins) so figures embed
    consistently across templates.
    """
    from docx.shared import Inches

    if not figures:
        return
    doc.add_heading("附录: 结果云图 / Figures", level=1)
    for name, path in figures.items():
        if not path.is_file():
            # Skip silently rather than crash export — a missing figure
            # was already surfaced by the renderer's stderr warning;
            # don't double-fail the report on its absence.
            continue
        doc.add_picture(str(path), width=Inches(6.0))
        cap = doc.add_paragraph(f"图 — {name}", style="Intense Quote")
        for run in cap.runs:
            run.italic = True


def export_docx(
    report: ReportSpec,
    bundle: EvidenceBundle,
    *,
    output_path: Path,
    template: "Optional[TemplateSpec]" = None,
    figures: "Optional[Dict[str, Path]]" = None,
    material: "Optional[Material]" = None,
) -> Path:
    """Materialise ``(report, bundle)`` to a DOCX at ``output_path``.

    Returns the resolved output path on success.

    When ``template`` is provided, the report is validated against the
    template contract via :func:`templates.validate_report` *before*
    any DOCX is written; a violation surfaces as
    :class:`TemplateValidationError`. ``template=None`` skips the
    template check (backward-compatible with W4 callers).

    Raises:
        ExportError: bundle linkage broken, an evidence_id cited in
            section content doesn't resolve in the bundle, or the
            output directory doesn't exist.
        TemplateValidationError: ``template`` was passed and the
            report violates its contract.
    """
    if template is not None:
        # Imported lazily to keep the module-import graph acyclic
        # (templates.py imports from app.models, exporter.py also
        # imports from app.models — fine — but we want to keep
        # templates.py optional from the exporter side).
        from app.services.report.templates import validate_report

        validate_report(report, bundle, template=template)

    if report.evidence_bundle_id != bundle.bundle_id:
        raise ExportError(
            f"bundle linkage broken: report.evidence_bundle_id="
            f"{report.evidence_bundle_id!r} but bundle.bundle_id="
            f"{bundle.bundle_id!r}"
        )

    # ADR-012 enforcement #1: every non-blank content line must cite EV-*.
    _check_every_content_line_cites_evidence(report.sections)

    # ADR-012 enforcement #2: every cited EV-* must resolve in the bundle.
    cited = find_cited_evidence_ids(report.sections)
    bundle_ids = {item.evidence_id for item in bundle.evidence_items}
    unresolved = cited - bundle_ids
    if unresolved:
        raise ExportError(
            f"section content cites evidence_id(s) {sorted(unresolved)!r} "
            f"that do not resolve in bundle {bundle.bundle_id!r} "
            f"(ADR-012). Bundle ids present: {sorted(bundle_ids)!r}."
        )

    # Path contract: resolve once (handles relative + ~), then require
    # the parent to exist *and be a directory*. A regular file at the
    # parent slot would otherwise leak ``NotADirectoryError`` from
    # ``doc.save``; we surface it as ``ExportError`` instead.
    resolved_output = output_path.expanduser().resolve()
    if not resolved_output.parent.is_dir():
        raise ExportError(
            f"output_path parent {resolved_output.parent} does not exist "
            "or is not a directory; create it before calling export_docx "
            "(the exporter does not mkdir on the caller's behalf)."
        )

    doc = Document()
    doc.add_heading(report.title, level=0)
    info_lines = [
        f"Report ID: {report.report_id}",
        f"Project ID: {report.project_id}",
        f"Template: {report.template_id}",
        f"Generated: {report.generated_at.isoformat()}",
        f"Evidence bundle: {report.evidence_bundle_id}",
    ]
    for ln in info_lines:
        doc.add_paragraph(ln)

    _render_section_tree(doc, report.sections)
    if material is not None:
        # W6a: § 材料属性 sits between the body sections and the
        # appendices so an engineer sweeping the document top-to-bottom
        # sees the material context before the evidence audit trail.
        # Order will be reshaped if W6e (模型概况) lands above it.
        _render_material_section(doc, material)
    _render_evidence_appendix(doc, bundle)
    if figures:
        _render_figures_appendix(doc, figures)

    doc.save(str(resolved_output))
    return resolved_output
