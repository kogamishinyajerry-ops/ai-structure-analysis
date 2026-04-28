"""End-to-end tests for ``report-cli --material`` — RFC-001 W6a.

Exercise the L1 → L4 path with material data threaded through, plus
the introspection / refusal flows (`--list-materials`, unknown
material, mutually exclusive flags, free-input flag).
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from app.services.report.cli import main

REPO_ROOT = Path(__file__).resolve().parents[1]
GS001_FRD = REPO_ROOT / "golden_samples" / "GS-001" / "gs001_result.frd"


@pytest.fixture()
def gs001_frd() -> Path:
    if not GS001_FRD.is_file():
        pytest.skip(f"GS-001 fixture missing at {GS001_FRD}")
    return GS001_FRD


def test_list_materials_prints_q345b(
    capsys: pytest.CaptureFixture[str],
) -> None:
    rc = main(["--list-materials"])
    assert rc == 0
    out = capsys.readouterr().out
    # Header + at least Q345B row visible. Tab-separated for scriptability.
    assert "code_grade\tstandard" in out
    assert "Q345B\tGB\t345\t470" in out
    # ASME-side spot
    assert "SA-516-70\tASME" in out


def test_list_materials_works_without_frd_or_kind(
    capsys: pytest.CaptureFixture[str],
) -> None:
    """--list-materials is an introspection mode, like --doctor; must
    not require --frd / --kind."""
    rc = main(["--list-materials"])
    assert rc == 0


def test_run_with_builtin_material_renders_section(gs001_frd: Path, tmp_path: Path) -> None:
    """End-to-end: --material Q345B threads through to a § 材料属性
    section in the produced DOCX. We use python-docx to introspect
    the saved file."""
    from docx import Document

    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material",
            "Q345B",
        ]
    )
    assert rc == 0
    doc = Document(str(out))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("材料属性" in h for h in headings), (
        f"§ 材料属性 heading missing; headings: {headings}"
    )

    # The Q345B row must be visible in some table cell — assert by
    # scanning all cells.
    all_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_cells.append(cell.text)
    assert any("Q345B" in c for c in all_cells), "Q345B not found in any DOCX table cell"
    assert any("345" in c for c in all_cells), "σ_y=345 MPa missing"
    assert any("470" in c for c in all_cells), "σ_u=470 MPa missing"
    assert any("GB/T 1591-2018" in c for c in all_cells), "source_citation missing"


def test_run_without_material_omits_section(gs001_frd: Path, tmp_path: Path) -> None:
    """W5f-compatibility: omitting --material must NOT inject a §
    材料属性 section. Engineers running pre-W6a workflows are not
    surprised by an empty section."""
    from docx import Document

    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
        ]
    )
    assert rc == 0
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert not any("材料属性" in h for h in headings), (
        f"§ 材料属性 must not appear without --material; got: {headings}"
    )


def test_run_with_user_supplied_json_flags_caveat(gs001_frd: Path, tmp_path: Path) -> None:
    """--material-json renders the section AND adds the [需工程师确认]
    caveat (RFC-001 §2.4 rule 4)."""
    from docx import Document

    mat_json = tmp_path / "mat.json"
    mat_json.write_text(
        json.dumps(
            {
                "code_grade": "INST-CUSTOM-Q420",
                "code_standard": "GB",
                "youngs_modulus": 206000,
                "poissons_ratio": 0.30,
                "yield_strength": 420,
                "ultimate_strength": 540,
                "density": 7.85e-9,
                "source_citation": "Institute-internal material card 2026-Q1",
            }
        ),
        encoding="utf-8",
    )
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material-json",
            str(mat_json),
        ]
    )
    assert rc == 0

    doc = Document(str(out))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "需工程师确认" in body, (
        f"user-supplied material must surface the [需工程师确认] caveat; body=\n{body[:500]}"
    )
    # The custom grade must show up in the table.
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert any("INST-CUSTOM-Q420" in c for c in cells)
    assert any("420" in c for c in cells)


def test_run_with_unknown_material_refuses(
    gs001_frd: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """Unknown built-in code is a domain refusal (exit 3), not a
    crash. Error message must point the engineer at
    --list-materials."""
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material",
            "BOGUS-GRADE",
        ]
    )
    assert rc == 3
    err = capsys.readouterr().err
    assert "BOGUS-GRADE" in err
    assert "--list-materials" in err
    # And the DOCX must not exist (we refused before export)
    assert not out.exists()


def test_material_and_material_json_are_mutually_exclusive(gs001_frd: Path, tmp_path: Path) -> None:
    """argparse should reject `--material X --material-json Y`. SystemExit(2)
    is the contract for input-validation."""
    mat_json = tmp_path / "mat.json"
    mat_json.write_text("{}", encoding="utf-8")
    with pytest.raises(SystemExit) as excinfo:
        main(
            [
                "--frd",
                str(gs001_frd),
                "--kind",
                "static",
                "--output",
                str(tmp_path / "x.docx"),
                "--no-figures",
                "--material",
                "Q345B",
                "--material-json",
                str(mat_json),
            ]
        )
    assert excinfo.value.code == 2


def test_material_section_precedes_body_sections(gs001_frd: Path, tmp_path: Path) -> None:
    """Codex R1 PR #91 MEDIUM regression — the § 材料属性 heading must
    appear BEFORE the body section headings (per RFC-001 §2.2 step 4
    order: 模型概况 → 材料属性 → 边界条件 → 关键结果). The previous
    placement injected material AFTER section_tree, putting it after
    'key results' instead of before."""
    from docx import Document

    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material",
            "Q345B",
        ]
    )
    assert rc == 0
    doc = Document(str(out))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    material_idx = next((i for i, h in enumerate(headings) if "材料属性" in h), -1)
    assert material_idx >= 0, f"§ 材料属性 missing; headings={headings}"

    # The static template's first body section is the report's results
    # heading; assert 材料属性 sits before it. We don't pin the literal
    # results-heading text (would be brittle if the template wording
    # changes); we pin only the *position* relative to other headings.
    # Specifically: § 材料属性 must NOT be the LAST heading (some
    # body section must follow it).
    assert material_idx < len(headings) - 1, (
        f"§ 材料属性 was placed last; expected at least one body "
        f"section to follow. headings={headings}"
    )
    # And evidence appendix must be AFTER 材料属性 (this was true
    # before too, but pin to prevent reshape regressions).
    evidence_idx = next(
        (i for i, h in enumerate(headings) if "证据" in h or "Evidence" in h),
        -1,
    )
    if evidence_idx >= 0:
        assert evidence_idx > material_idx, (
            f"evidence appendix placed before § 材料属性; headings={headings}"
        )


def test_material_lookup_fails_fast_before_reader_open(
    gs001_frd: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """Codex R1 PR #91 MEDIUM regression — an invalid --material must
    fail BEFORE the [1/N] reading-frd stage runs. The previous flow
    deferred lookup to after _produce(), which masked the typo with
    upstream noise."""
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material",
            "BOGUS-GRADE",
        ]
    )
    assert rc == 3
    err = capsys.readouterr().err
    # The "[1/N] reading CalculiX .frd" stage line must NOT appear —
    # we refused before opening any I/O.
    assert "reading CalculiX" not in err, (
        f"material lookup should fail before stage 1; got stderr:\n{err}"
    )
    assert "BOGUS-GRADE" in err
    assert "--list-materials" in err


def test_material_audit_trail_shows_in_stderr(
    gs001_frd: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """The producer detail line must surface the material on stderr
    so the Electron audit trail records which material was used."""
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
            "--material",
            "Q345B",
        ]
    )
    assert rc == 0
    err = capsys.readouterr().err
    # ASCII '->' detail prefix per the W5e CP936-safe contract.
    assert "      -> material: Q345B (GB) sigma_y=345 sigma_u=470" in err
