"""Tests for ``report-cli`` with figure rendering — RFC-001 W5f.

Exercises the L1-direct → viz path that materialises three PNG figures
alongside the DOCX and embeds them as a Figures appendix. Skipped if
pyvista or vtk aren't importable in the test environment (these deps
are heavy and may not be present in trimmed CI variants).
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest

# pyvista is a hard dep declared in pyproject.toml, but not every test
# environment has the GL stack pyvista needs to off-screen render. Skip
# the whole module if either is missing — the figure feature is opt-out
# at runtime via --no-figures, so absence shouldn't fail the suite.
pytest.importorskip("pyvista")
pytest.importorskip("vtk")

from app.services.report.cli import main  # noqa: E402

REPO_ROOT = Path(__file__).resolve().parents[1]
GS001_FRD = REPO_ROOT / "golden_samples" / "GS-001" / "gs001_result.frd"

# pyvista off-screen mode is the only safe default for headless tests.
os.environ.setdefault("PYVISTA_OFF_SCREEN", "true")


@pytest.fixture()
def gs001_frd() -> Path:
    if not GS001_FRD.is_file():
        pytest.skip(f"GS-001 fixture missing at {GS001_FRD}")
    return GS001_FRD


def test_default_run_renders_three_figures(
    gs001_frd: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """A default run (no --no-figures) writes mesh / displacement /
    von_mises PNGs into <output>.figs/ and embeds them in the DOCX."""
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
        ]
    )
    assert rc == 0

    figs_dir = tmp_path / "report.figs"
    assert figs_dir.is_dir()
    expected = {"mesh.png", "displacement.png", "von_mises.png"}
    actual = {p.name for p in figs_dir.iterdir() if p.suffix == ".png"}
    assert expected.issubset(actual), f"missing figures: {expected - actual}"

    # Each PNG must be a real, non-trivial render — not a placeholder.
    for name in expected:
        size = (figs_dir / name).stat().st_size
        assert size > 5_000, f"{name} is {size}B — looks like a stub render"

    # The Electron-side gallery hook depends on stderr "figure: <path>"
    # lines — verify each is announced exactly once.
    err = capsys.readouterr().err
    for name in expected:
        path_str = str(figs_dir / name)
        assert err.count(f"figure: {path_str}") == 1, (
            f"figure announcement for {name} not found exactly once;\nstderr was:\n{err}"
        )

    # The "[3/5] rendering figures" stage prefix must appear too.
    assert "[3/5] rendering figures:" in err


def test_no_figures_flag_skips_rendering(
    gs001_frd: Path,
    tmp_path: Path,
    capsys: pytest.CaptureFixture[str],
) -> None:
    """--no-figures must suppress the figs/ directory AND drop the
    rendering stage from the [N/T] denominator."""
    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
        ]
    )
    assert rc == 0

    figs_dir = tmp_path / "report.figs"
    assert not figs_dir.exists()

    err = capsys.readouterr().err
    assert "rendering figures:" not in err
    assert "figure:" not in err
    # 4-stage denominator (no figures).
    assert "[1/4]" in err
    assert "[4/4] exporting DOCX:" in err


def test_figures_dir_override(
    gs001_frd: Path,
    tmp_path: Path,
) -> None:
    """--figures-dir overrides the default <output>.figs/ location."""
    out = tmp_path / "report.docx"
    custom = tmp_path / "custom_figs"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--figures-dir",
            str(custom),
        ]
    )
    assert rc == 0
    assert custom.is_dir()
    assert (custom / "mesh.png").is_file()
    assert (custom / "displacement.png").is_file()
    assert (custom / "von_mises.png").is_file()
    # The default <output>.figs path must NOT have been created.
    assert not (tmp_path / "report.figs").exists()


def test_docx_embeds_figures(
    gs001_frd: Path,
    tmp_path: Path,
) -> None:
    """The exported DOCX must include the three figures as embedded
    images. python-docx's API exposes inline_shapes; we just count."""
    from docx import Document

    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
        ]
    )
    assert rc == 0

    doc = Document(str(out))
    # 3 figures = 3 inline shapes (mesh / displacement / von_mises).
    assert len(doc.inline_shapes) == 3, (
        f"expected 3 embedded figures, found {len(doc.inline_shapes)}"
    )
    # Headings should now include the Figures appendix label.
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Figures" in h or "云图" in h for h in headings), (
        f"Figures appendix heading missing; headings: {headings}"
    )


def test_no_figures_docx_has_no_embedded_images(
    gs001_frd: Path,
    tmp_path: Path,
) -> None:
    """With --no-figures, the DOCX must not have an inline image."""
    from docx import Document

    out = tmp_path / "report.docx"
    rc = main(
        [
            "--frd",
            str(gs001_frd),
            "--kind",
            "static",
            "--output",
            str(out),
            "--no-figures",
        ]
    )
    assert rc == 0
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 0
